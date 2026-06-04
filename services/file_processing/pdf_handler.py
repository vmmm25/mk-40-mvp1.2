import logging
from pathlib import Path
from .utils import get_output_path, gemini_client, get_file_size_str

logger = logging.getLogger(__name__)

def process_pdf(path: Path, action: str, params: dict, speak=None) -> str:
    action = action or "summarize"

    def _extract_pdf_text(max_chars=50000) -> str:
        text = ""
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    text += (page.extract_text() or "") + "\n"
        except ImportError:
            try:
                import PyPDF2
                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
            except ImportError:
                return ""
        return text[:max_chars]

    if action in ("summarize", "extract_text", "translate_hint", "analyze", "reformat"):
        text = _extract_pdf_text()
        if not text.strip():
            return "Could not extract text from PDF (may be scanned/image-based)."

        if action == "extract_text":
            out = get_output_path(path, "text", ".txt")
            out.write_text(text, encoding="utf-8")
            return f"Text extracted ({len(text)} chars). Saved: {out.name}"

        prompt_map = {
            "summarize":      f"Summarize this PDF document concisely:\n\n{text}",
            "analyze":        f"Analyze this document thoroughly:\n\n{text}",
            "translate_hint": f"What language is this document in and what does it say? Summarize:\n\n{text}",
            "reformat":       f"Reformat this text cleanly with proper structure:\n\n{text}",
        }
        try:
            model    = gemini_client()
            response = model.generate_content(prompt_map.get(action, f"Analyze:\n\n{text}"))
            result   = response.text.strip()
            if len(result) > 600 and params.get("save", True):
                out = get_output_path(path, action, ".txt")
                out.write_text(result, encoding="utf-8")
                return f"{result[:400]}...\n\nFull result saved: {out.name}"
            return result
        except Exception as e:
            logger.exception("AI analysis failed")
            return f"AI analysis failed: {e}"

    if action == "info":
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                pages = len(pdf.pages)
            return f"PDF: {pages} pages, size: {get_file_size_str(path)}"
        except Exception:
            return f"PDF size: {get_file_size_str(path)}"

    if action == "to_word":
        text = _extract_pdf_text()
        if not text:
            return "Could not extract text to convert."
        try:
            from docx import Document
            doc  = Document()
            doc.add_heading(path.stem, 0)
            for para in text.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
            out = get_output_path(path, "converted", ".docx")
            doc.save(out)
            return f"Converted to Word document. Saved: {out.name}"
        except ImportError:
            return "python-docx not installed. Run: pip install python-docx"

    return f"Unknown PDF action: '{action}'. Try: summarize, extract_text, info, to_word"
